# coding:utf-8

import os
from docx import Document
from docx.shared import Pt
import logging
import time
from wxpay import *
from flask import Flask, request, jsonify, current_app
import socket

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False

# 获取本机计算机名称s
hostname = socket.gethostname()
socket.gethostbyname(hostname)
app = Flask(__name__)
appid = 'wxbcdbe97d7b353c80'  # 小程序ID
mch_id = '1526671931'  # 商户号
total_fee = '5000'  # 总金额
attach = 'test'
# 获取本机ip
spbill_create_ip = socket.gethostbyname(hostname)  # 终端IP
notify_url = 'https://file.cumtlee.cn/wxpay/notify'  # 通知地址
trade_type = 'JSAPI'  # 交易类型
merchant_key = 'F061CF04D8335A953ECFB982444307DD'  # 商户KEY


@app.route('/')
def index():
    '''
    :return:
    '''
    return jsonify({'errcode': 0, 'errmsg': 'ok'})


@app.route('/wxpay/pay')
def create_pay():
    '''
    请求支付
    :return:
    '''
    data = {
        'appid': appid,
        'mch_id': mch_id,
        'attach': attach,
        'nonce_str': get_nonce_str(),
        'body': 'JSAPI-Pay',  # 商品描述
        'out_trade_no': str(int(time.time())),  # 商户订单号
        'total_fee': total_fee,
        'spbill_create_ip': spbill_create_ip,
        'notify_url': notify_url,
        'trade_type': trade_type,
        'openid': request.args.get('openid')
    }
    '''
    生成签名
    :return:
    '''
    # print(data)
    wxpay = WxPay(merchant_key, data)
    pay_info = wxpay.get_pay_info()
    if pay_info:
        return jsonify(pay_info)
    return jsonify({'errcode': 40001, 'errmsg': '请求支付失败'})


@app.route('/wxpay/notify', methods=['POST'])
def wxpay():
    '''
    支付回调通知
    :return:
    '''
    if request.method == 'POST':
        logging.info(xml_to_dict(request.data))
        result_data = {
            'return_code': 'SUCCESS',
            'return_msg': 'OK'
        }
        return dict_to_xml(result_data), {'Content-Type': 'application/xml'}


@app.route("/downfile/<filename>")
def getfile(filename):
    '''
    生成准考证
    :return:
    '''
    param = request.args
    # print(param)
    document = Document(r"zkz.docx")
    tables = document.tables
    arr = [
        [param.get('stu_name'), param.get('exam_time')],
        [param.get('stu_college'), param.get('exam_site')],
        [param.get('stu_id'), param.get('exam_location')],
        [param.get('stu_school'), param.get('exam_seat_num')],
        [param.get('stu_catetory'), param.get('exam_zkz_num')]
    ]
    for i in range(5):
        tables[0].cell(i + 1, 1).paragraphs[0].add_run(arr[i][0]).font.size = Pt(14)
        tables[0].cell(i + 1, 3).paragraphs[0].add_run(arr[i][1]).font.size = Pt(14)

    path = './files/' + param.get('stu_id') + '.docx'
    document.save(path)

    def generate():
        with open(path, 'rb') as f:
            yield from f
        os.remove(path)

    r = current_app.response_class(generate(),
                                   mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    exam_zkz_num = param.get('exam_zkz_num') + '.docx'
    r.headers.set('Content-Disposition', 'attachment', filename=exam_zkz_num)
    return r


if __name__ == '__main__':
    app.run(debug=True)
